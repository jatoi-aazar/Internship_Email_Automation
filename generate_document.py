from datetime import datetime
from docx import Document
from config import TEMPLATE_PATH
from datetime import timedelta


def generate_offer_letter(applicant):
    
    document = Document(TEMPLATE_PATH)

    today = datetime.today()

    start_date = today + timedelta(days=3)
    
    duration = "30 Days"

    end_date = start_date + timedelta(days=30)

    DATE = today.strftime('%d %B %Y')
    START_DATE = start_date.strftime('%d %B %Y')
    END_DATE = end_date.strftime('%d %B %Y')

    replace_with = {
        "{{DATE}}": DATE,
        "{{NAME}}": applicant["name"],
        "{{DEPARTMENT}}": applicant["department"],
        "{{START_DATE}}": START_DATE,
        "{{DURATION}}": duration,
        "{{END_DATE}}": END_DATE,
        "{{ROLE}}": f'{applicant["department"] } Intern',
        "{{WORKING_HOURS}}" : "9:00 AM - 6:00 PM",
        "{{WORK_LOCATION}}": "Remote",
        "{{HR_NAME}}": "Aazar Jatoi",
        "{{HR_DESIGNATION}}": "Python Developer"
    }

    for paragraph in  document.paragraphs:  
        replace_placeholders(paragraph , replace_with)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders(paragraph,  replace_with)
                    
    document_path = save_document(document , applicant["name"])
    return document_path

def replace_placeholders(paragraph , replace_with):
        final_text = paragraph.text
        for placeholder , value in replace_with.items():
            final_text = final_text.replace(placeholder , str(value))
        paragraph.text = final_text    

 
def save_document(document , name):
    output_path = f"generated_pdfs/{name}_offer_letter.docx"
    document.save(output_path)

    return output_path

